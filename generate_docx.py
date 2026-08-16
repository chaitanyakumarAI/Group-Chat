import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_screenshot_box(doc, title_text, instruction_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_background(cell, "F8FAFC")
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    
    r1 = p.add_run("📸 " + title_text + "\n")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    
    r2 = p.add_run(instruction_text + "\n\n")
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    r3 = p.add_run("[ Paste Screenshot Here ]")
    r3.font.italic = True
    r3.font.bold = True
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(4)

def create_docx_report():
    doc_path = os.path.join(os.path.dirname(__file__), "Group_Chat_Architecture_Report.docx")
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Document Header
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("Real-Time Group Chat Application")
    r_title.font.name = "Arial"
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("Computer System Design (CSD) — Assignment 4 | Persistence, Security & Verification Report")
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    doc.add_paragraph()

    # Quick Info Box Table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        [("Course", "Computer System Design (CSD)"), ("Assignment", "Assignment 4")],
        [("Host SSH Machine", "student@10.1.75.51 (Port 2237)"), ("Internal Server Port", "5000")],
        [("Public TA Testing URL", "http://10.1.75.51:5237/"), ("Tech Stack", "Flask, Flask-SocketIO, SQLite, AES-256-GCM, Ed25519")]
    ]

    for r_idx, row in enumerate(info_table.rows):
        for c_idx, cell in enumerate(row.cells):
            set_cell_background(cell, "F1F5F9")
            label, val = info_data[r_idx][c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r1 = p.add_run(f"{label}: ")
            r1.font.bold = True
            r1.font.size = Pt(9.5)
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            if label == "Public TA Testing URL":
                r2.font.bold = True
                r2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph()

    # Section 1: Executive Summary
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Executive Summary & Lab Network Mapping")
    r_h1.font.bold = True
    r_h1.font.size = Pt(13)
    r_h1.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.15
    r_ex = p_exec.add_run(
        "This project implements a multi-user, real-time Group Chat Application for Computer System Design (CSD) Assignment 4. "
        "Built using WebSockets (Flask-SocketIO) and SQLite persistence, it incorporates end-to-end security layers "
        "(AES-256-GCM confidentiality and Ed25519 digital signatures). The backend server runs on the allotted SSH laboratory machine (student@10.1.75.51).\n\n"
        "Network & Port Forwarding Details:\n"
        "The Flask backend process binds internally to port 5000. The laboratory network NAT router maps SSH port 2237 to external web port 5237. "
        "Therefore, the official public testing URL accessible to all 4 student clients and TAs is: "
    )
    r_ex.font.size = Pt(9.5)
    r_url = p_exec.add_run("http://10.1.75.51:5237/")
    r_url.font.bold = True
    r_url.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    r_url.font.size = Pt(9.5)

    doc.add_paragraph()

    # Section 2: Team Members & Contributions
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. Team Members & Individual Contributions")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13)
    r_h2.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    team_table = doc.add_table(rows=5, cols=5)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Role / Designation", "Student Member Name", "Roll Number", "SSH Machine Command", "Key Technical Contribution"]
    
    # Format Headers
    hdr_cells = team_table.rows[0].cells
    for i, title in enumerate(headers):
        set_cell_background(hdr_cells[i], "1E1B4B")
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    members_data = [
        ("Group Head (Host)", "Ranga Chandra Naga Venkata Chaitanya Kumar", "12341740", "ssh -p 2237 student@10.1.75.51", "Backend System Architecture, Flask-SocketIO Core, AES-256-GCM & Ed25519 Cryptographic Pipeline, SSH Server Deployment & Daemon Management."),
        ("Member 2", "Bhukya Raju", "12340520", "ssh -p 2238 student@10.1.75.51", "Front-End UI Design, Cyberpunk Aurora Glassmorphism Theme, Socket.IO Client Event Listeners, Dynamic Gradient Avatar Generator."),
        ("Member 3", "V.G.N. Harshitha", "12342310", "ssh -p 2239 student@10.1.75.51", "Database Schema Architecture (chat.db), SQLite Message Persistence Engine, Reconnection State Management & History Pre-loading."),
        ("Member 4", "Maloth Madhu", "12341370", "ssh -p 2240 student@10.1.75.51", "Multi-Client Integration & Testing across SSH Machines (10.1.75.51:5237), Graceful Disconnection Handling, SQLite Inspection Scripts.")
    ]

    for r_idx, row in enumerate(members_data):
        row_cells = team_table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            p = row_cells[c_idx].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if c_idx == 0:
                r.font.bold = True

    doc.add_paragraph()

    # Section 3: Key Technical Demonstrations & Screenshots
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Demonstrations: Persistence, Tamper Detection & Signatures")
    r_h3.font.bold = True
    r_h3.font.size = Pt(13)
    r_h3.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    add_screenshot_box(
        doc,
        "Screenshot 1: Real-Time Multi-User Chat UI Demonstration",
        "Demonstrates active chat on http://10.1.75.51:5237/ with multiple client connections, dynamic gradient avatars, and live typing status."
    )

    add_screenshot_box(
        doc,
        "Screenshot 2: Demonstration of Message Persistence (Server Restart)",
        "Demonstrates persistence: The python3 server.py process was terminated and restarted. Upon client reconnection, all historical messages were automatically loaded from SQLite (chat.db)."
    )

    add_screenshot_box(
        doc,
        "Screenshot 3: Demonstration of Tamper Detection (AES-256-GCM Auth Tag Failure)",
        "Demonstrates Tamper Detection: After corrupting 1 byte of stored ciphertext in chat.db using demo_tamper.py, AES-256-GCM authentication tag verification failed, displaying [TAMPER DETECTED: ciphertext/authentication tag invalid]."
    )

    add_screenshot_box(
        doc,
        "Screenshot 4: Demonstration of Signature Verification (Ed25519)",
        "Demonstrates Digital Signatures: Terminal output of sqlite3 ~/chat_app/chat.db showing per-user Ed25519 public keys in signing_keys table and raw cryptographic signatures in messages table."
    )

    doc.add_paragraph()

    # Section 4: Requirement Verification Table
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. Core Features & Requirement Verification Matrix")
    r_h4.font.bold = True
    r_h4.font.size = Pt(13)
    r_h4.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    req_table = doc.add_table(rows=9, cols=3)
    req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    req_headers = ["Requirement / Feature", "Implementation Mechanism", "Verification Status"]
    for i, title in enumerate(req_headers):
        set_cell_background(req_table.rows[0].cells[i], "4F46E5")
        p = req_table.rows[0].cells[i].paragraphs[0]
        r = p.add_run(title)
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)

    reqs = [
        ("Real-Time Message Broadcast", "Bi-directional WebSocket event broadcasting via Flask-SocketIO", "VERIFIED"),
        ("User Join/Leave Alerts", "Automatic broadcast of user arrival/departure pills to room", "VERIFIED"),
        ("User Identification", "Unique username validation & dynamic gradient avatar generation", "VERIFIED"),
        ("Client Disconnection Cleanup", "Server socket connection monitoring & cleanup on tab close", "VERIFIED"),
        ("SQLite Message Persistence", "All messages saved in chat.db; auto-loaded upon joining chat", "VERIFIED"),
        ("Tamper Detection (AES-GCM)", "Payload encrypted with room key prior to database insertion; auth tag verified on read", "VERIFIED"),
        ("Digital Signatures (Ed25519)", "Cryptographic digital signature per user stored & verified against public key", "VERIFIED"),
        ("Public TA Access URL", "Mapped external port forwarded to http://10.1.75.51:5237/", "VERIFIED")
    ]

    for r_idx, (req, impl, status) in enumerate(reqs):
        cells = req_table.rows[r_idx + 1].cells
        p0 = cells[0].paragraphs[0]
        p0.add_run(req).font.size = Pt(8.5)
        p1 = cells[1].paragraphs[0]
        p1.add_run(impl).font.size = Pt(8.5)
        p2 = cells[2].paragraphs[0]
        r_stat = p2.add_run(status)
        r_stat.font.size = Pt(8.5)
        r_stat.font.bold = True
        r_stat.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)

    doc.add_paragraph()

    # Section 5: Commands to Reproduce & Verify Screenshots
    h5 = doc.add_paragraph()
    r_h5 = h5.add_run("5. Commands to Reproduce & Verify Tests on SSH Server")
    r_h5.font.bold = True
    r_h5.font.size = Pt(13)
    r_h5.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    p_cmd = doc.add_paragraph()
    r_c1 = p_cmd.add_run("1. To inspect stored encrypted messages & Ed25519 signatures:\n")
    r_c1.font.size = Pt(9)
    r_c2 = p_cmd.add_run("sqlite3 ~/chat_app/chat.db \"SELECT id, sender, ciphertext, signature, timestamp FROM messages;\"\n\n")
    r_c2.font.name = "Consolas"
    r_c2.font.size = Pt(8.5)

    r_c3 = p_cmd.add_run("2. To verify user Ed25519 public keys:\n")
    r_c3.font.size = Pt(9)
    r_c4 = p_cmd.add_run("sqlite3 ~/chat_app/chat.db \"SELECT username, public_key FROM signing_keys;\"\n\n")
    r_c4.font.name = "Consolas"
    r_c4.font.size = Pt(8.5)

    r_c5 = p_cmd.add_run("3. To trigger Tamper Detection for Screenshot 3:\n")
    r_c5.font.size = Pt(9)
    r_c6 = p_cmd.add_run("python3 ~/chat_app/demo_tamper.py")
    r_c6.font.name = "Consolas"
    r_c6.font.size = Pt(8.5)

    doc.add_paragraph()

    # Section 6: Submission Links
    h6 = doc.add_paragraph()
    r_h6 = h6.add_run("6. Official Submission Links")
    r_h6.font.bold = True
    r_h6.font.size = Pt(13)
    r_h6.font.color.rgb = RGBColor(0x1E, 0x1B, 0x4B)

    p_links = doc.add_paragraph()
    p_links.add_run("• GitHub Repository: ").font.size = Pt(9.5)
    r_gh = p_links.add_run("https://github.com/chaitanyakumarAI/Group-Chat\n")
    r_gh.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    r_gh.font.size = Pt(9.5)

    p_links.add_run("• Official Testing Public URL: ").font.size = Pt(9.5)
    r_url2 = p_links.add_run("http://10.1.75.51:5237/\n")
    r_url2.font.bold = True
    r_url2.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
    r_url2.font.size = Pt(9.5)

    p_links.add_run("• Host SSH Server Command: ").font.size = Pt(9.5)
    r_ssh = p_links.add_run("ssh -p 2237 student@10.1.75.51")
    r_ssh.font.name = "Consolas"
    r_ssh.font.size = Pt(9)

    doc.save(doc_path)
    print(f"Word Document successfully generated at: {doc_path}")

if __name__ == "__main__":
    create_docx_report()
